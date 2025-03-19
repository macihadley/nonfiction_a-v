from vaderSentiment.vaderSentiment import SentimentIntensityAnalyzer
import nltk
import os
from tkinter import Tk, filedialog
from docx import Document
from docx.shared import RGBColor

# Download required NLTK data (if needed)
nltk.download('stopwords')
nltk.download('punkt')
nltk.download('wordnet')

# Function to analyze and highlight sentiment in a DOCX file
def analyze_sentiment_vader(text, doc):
    sid = SentimentIntensityAnalyzer()
    words = text.split()

    paragraph = doc.add_paragraph()
    for word in words:
        sentiment = sid.polarity_scores(word)['compound']
        run = paragraph.add_run(word + " ")

        # Assign color based on sentiment polarity
        if sentiment > 0:
            run.font.color.rgb = RGBColor(255, 140, 0)  # Dark Orange
        elif sentiment < 0:
            run.font.color.rgb = RGBColor(0, 0, 139)  # Dark Blue
        else:
            run.font.color.rgb = RGBColor(0, 0, 0)  # Black

# Function to process all text files in a folder and save results as DOCX
def process_folder(input_folder, output_folder):
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    for filename in os.listdir(input_folder):
        if filename.endswith(".txt"):
            file_path = os.path.join(input_folder, filename)
            output_file_path = os.path.join(output_folder, f"{os.path.splitext(filename)[0]}_sentiment.docx")

            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    lines = file.readlines()

                title = lines[0].strip() if len(lines) > 0 else "Untitled"
                author = lines[1].strip() if len(lines) > 1 else "Unknown"
                body = ' '.join(lines[2:])

                # Create DOCX document
                doc = Document()
                doc.add_heading(title, level=1)
                #doc.add_paragraph(f"Author: {author}\n")

                analyze_sentiment_vader(body, doc)

                # Save the document
                doc.save(output_file_path)
                print(f"Analysis saved to {output_file_path}")

            except Exception as e:
                print(f"Error processing {filename}: {e}")

if __name__ == "__main__":
    # Use Tkinter to select folders dynamically
    Tk().withdraw()  # Hide the root Tkinter window
    input_folder = filedialog.askdirectory(title="Select Folder with TXT Files")
    output_folder = filedialog.askdirectory(title="Select Output Folder")

    if input_folder and output_folder:
        print(f"Processing files in: {input_folder}")
        print(f"Saving DOCX results in: {output_folder}")
        process_folder(input_folder, output_folder)
    else:
        print("No folder selected. Exiting.")

